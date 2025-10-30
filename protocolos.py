from io import BytesIO
from pathlib import Path
from typing import Dict, List

import streamlit as st

try:
    from docx import Document  # type: ignore
except ImportError:  # pragma: no cover - handled at runtime
    Document = None  # type: ignore

try:
    import PyPDF2  # type: ignore
except ImportError:  # pragma: no cover - handled at runtime
    PyPDF2 = None  # type: ignore


def extract_text(file_buffer: BytesIO, extension: str) -> str:
    """Retorna texto simples a partir de um DOCX ou PDF."""
    extension = extension.lower()
    if extension == "docx":
        if Document is None:
            raise ImportError("python-docx nao esta instalado.")
        file_buffer.seek(0)
        document = Document(file_buffer)
        paragraphs: List[str] = [paragraph.text for paragraph in document.paragraphs]

        # Tambem percorre celulas de tabela para nao perder texto.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        return "\n".join(filter(None, paragraphs))

    if extension == "pdf":
        if PyPDF2 is None:
            raise ImportError("PyPDF2 nao esta instalado.")
        file_buffer.seek(0)
        reader = PyPDF2.PdfReader(file_buffer)
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        return "\n".join(pages)

    raise ValueError("Tipo de arquivo nao suportado.")


def parse_key_values(raw_text: str) -> Dict[str, str]:
    """Transforma linhas no formato chave: valor em um dicionario."""
    pairs: Dict[str, str] = {}
    for line in raw_text.splitlines():
        cleaned = line.strip()
        if not cleaned or cleaned.startswith("#"):
            continue
        if ":" not in cleaned:
            continue
        key, value = cleaned.split(":", 1)
        key = key.strip()
        value = value.strip()
        if key:
            pairs[key] = value
    return pairs


def _replace_in_paragraphs(paragraphs, values: Dict[str, str]) -> None:
    """Substitui placeholders {{chave}} preservando estilo basico."""
    for paragraph in paragraphs:
        original = paragraph.text
        updated = original
        for key, value in values.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in updated:
                updated = updated.replace(placeholder, value)
        if updated != original:
            paragraph.text = updated


def fill_template_with_values(template_bytes: BytesIO, values: Dict[str, str]) -> BytesIO:
    """Aplica os valores encontrados em um template DOCX e retorna bytes atualizados."""
    if Document is None:
        raise ImportError("python-docx nao esta instalado.")

    template_bytes.seek(0)
    document = Document(template_bytes)

    _replace_in_paragraphs(document.paragraphs, values)

    for section in document.sections:
        if section.header:
            _replace_in_paragraphs(section.header.paragraphs, values)
        if section.footer:
            _replace_in_paragraphs(section.footer.paragraphs, values)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, values)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


st.set_page_config(
    page_title="Preenchimento de Template",
    page_icon=":memo:",
    layout="centered",
)

st.title("Gerador de documentos a partir de Template")
st.markdown(
    """
Envie o documento **template** (com placeholders no formato `{{chave}}`) e o
documento **base** contendo os dados no formato `chave: valor`, um por linha.
Suporta arquivos `.docx` ou `.pdf` para extracao de texto. A aplicacao monta um
novo DOCX preenchido quando o template for DOCX.
"""
)

col_template, col_base = st.columns(2)
with col_template:
    template_upload = st.file_uploader(
        "Selecione o template (.docx ou .pdf)",
        type=["docx", "pdf"],
        key="template_uploader",
    )
with col_base:
    base_upload = st.file_uploader(
        "Selecione o documento base (.docx ou .pdf)",
        type=["docx", "pdf"],
        key="base_uploader",
    )

if template_upload and base_upload:
    template_extension = Path(template_upload.name).suffix.lstrip(".").lower()
    base_extension = Path(base_upload.name).suffix.lstrip(".").lower()

    template_bytes_raw = template_upload.getvalue()
    base_bytes_raw = base_upload.getvalue()

    st.info(
        f"Template: **{template_upload.name}** ({template_extension.upper()}) | "
        f"Base: **{base_upload.name}** ({base_extension.upper()})"
    )

    try:
        base_text = extract_text(BytesIO(base_bytes_raw), base_extension)
    except ImportError as missing_dep:
        st.error(
            f"Nao foi possivel ler o documento base: {missing_dep}. "
            "Instale a dependencia e recarregue a pagina."
        )
        base_text = ""
    except Exception as exc:  # pragma: no cover - feedback ao usuario
        st.error(f"Erro ao ler o documento base: {exc}")
        base_text = ""

    if base_text:
        st.subheader("Pre-visualizacao dos dados extraidos")
        st.text_area(
            label="Dados detectados",
            value=base_text,
            height=200,
            help="Revise o texto e confirme se as linhas estao no formato chave: valor.",
        )

    values = parse_key_values(base_text)
    if not values:
        st.warning(
            "Nao foram encontrados pares `chave: valor` no documento base. "
            "Verifique o arquivo enviado."
        )

    filled_template: BytesIO | None = None
    if values and template_extension == "docx":
        try:
            filled_template = fill_template_with_values(
                BytesIO(template_bytes_raw), values
            )
        except ImportError as missing_dep:
            st.error(
                f"Nao foi possivel editar o template DOCX: {missing_dep}. "
                "Instale a dependencia e recarregue a pagina."
            )
        except Exception as exc:  # pragma: no cover - feedback ao usuario
            st.error(f"Ocorreu um erro ao preencher o template: {exc}")

    if template_extension != "docx":
        st.info(
            "No momento apenas templates DOCX podem ser alterados automaticamente. "
            "Voce ainda pode usar o texto extraido para atualizar o PDF manualmente."
        )

    if filled_template:
        st.success("Template preenchido com sucesso! Faca o download abaixo.")
        st.download_button(
            label="Baixar documento preenchido",
            data=filled_template.getvalue(),
            file_name="template-preenchido.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
